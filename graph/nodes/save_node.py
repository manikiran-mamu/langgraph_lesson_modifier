# graph/nodes/save_node.py

from graph.schema import State
from docx import Document
from docx.shared import Pt
import os
import uuid

TEMPLATE_PATH = "templates/lesson_template.docx"
OUTPUT_DIR = "data/outputs/word"

# Mapping for the activity table
ACTIVITY_TABLE_MAPPING = {
    ("Lesson Intro", "Teacher Activities"): "intro_teacher",
    ("Demonstration, model, or mini-lesson (“I DO”)", "Teacher Activities"): "i_do_teacher",
    ("Shared Learning or Guided Practice (“We Do”)", "Teacher Activities"): "we_do_teacher",
    ("Independent or Collaborative Small Group Work (“You Do”) Assessment)", "Teacher Activities"): "you_do_teacher",
    ("Lesson Intro", "Desired Student Actions and Potential Misconceptions"): "intro_student",
    ("Demonstration, model, or mini-lesson (“I DO”)", "Desired Student Actions and Potential Misconceptions"): "i_do_student",
    ("Shared Learning or Guided Practice (“We Do”)", "Desired Student Actions and Potential Misconceptions"): "we_do_student",
    ("Independent or Collaborative Small Group Work (“You Do”) Assessment)", "Desired Student Actions and Potential Misconceptions"): "you_do_student",
}

def insert_into_cell(cell, content: str):
    """Insert formatted text into a Word cell with styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(content.strip())
    run.font.name = "Poppins"
    run.font.size = Pt(11)

def save_node(state: State) -> State:
    print("📝 Filling lesson template...")

    sections = state.get("sections")
    if not sections:
        raise ValueError("No 'sections' data found in state. Did generate_node run correctly?")

    doc = Document(TEMPLATE_PATH)

    # Replace Title (first paragraph)
    if doc.paragraphs:
        title_paragraph = doc.paragraphs[0]
        title_paragraph.text = "Lesson Plan"
        if title_paragraph.runs:
            run = title_paragraph.runs[0]
        else:
            run = title_paragraph.add_run()
        run.font.name = "Poppins"
        run.font.size = Pt(24)
        run.bold = True

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # Check if it's the Standards/Content/Language/Purpose table
        if "Standards Addressed" in rows[0].cells[0].text and "Objectives or Essential Question" in rows[0].cells[1].text:
            try:
                standards_cell = rows[1].cells[0]
                content_cell = rows[1].cells[1]
                purpose_cell = rows[2].cells[0]
                language_cell = rows[1].cells[1]

                # Fill standards
                if "standards" in sections:
                    insert_into_cell(standards_cell, sections["standards"])

                # Fill content (look for "CONTENT:" block)
                if "content" in sections and "CONTENT:" in content_cell.text:
                    insert_into_cell(content_cell, sections["content"])

                # Fill language (look for "LANGUAGE:" block)
                if "language" in sections and "LANGUAGE:" in content_cell.text:
                    insert_into_cell(language_cell, sections["language"])

                # Fill purpose
                if "purpose" in sections:
                    insert_into_cell(purpose_cell, sections["purpose"])
            except Exception as e:
                print("⚠️ Could not fill Standards table:", str(e))
            continue

        # Handle activity table (student/teacher activities)
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        for row in rows[1:]:
            row_label = row.cells[0].text.strip()
            for col_idx, col_text in enumerate(header_cells[1:], start=1):
                key = (row_label, col_text.strip())
                if key in ACTIVITY_TABLE_MAPPING:
                    section_key = ACTIVITY_TABLE_MAPPING[key]
                    if section_key in sections:
                        content = sections[section_key]
                        if isinstance(content, list):
                            content = "\n".join(f"• {item}" for item in content)
                        insert_into_cell(row.cells[col_idx], content)

    # Save filled document
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = f"lesson_plan_filled_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)

    print(f"✅ Lesson plan saved at: {output_path}")
    return state.update({"final_output_docx": output_path})