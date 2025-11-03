import os
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List

OUTPUT_DIR = "data/outputs/source_materials"

def save_source_material_doc(processed_paragraphs: List[str]) -> str:
    """
    Save processed lesson paragraphs into a formatted Word document.
    Title: 'Source Material Text' (Poppins, 24pt Bold)
    Paragraphs: Poppins, 14pt, separated by line breaks
    """
    doc = Document()

    # 🟦 Title
    title = doc.add_paragraph()
    run = title.add_run("Source Material Text")
    run.font.name = "Poppins"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # 📘 Paragraphs
    for para in processed_paragraphs:
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.name = "Poppins"
        run.font.size = Pt(12)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph()  # Blank line between paragraphs

    # 💾 Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = f"source_material_{uuid.uuid4().hex}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)

    print(f"✅ Source material document saved at: {path}")
    return path