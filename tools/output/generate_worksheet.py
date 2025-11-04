import os
import re
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



def add_horizontal_line(doc):
    """Add a true continuous horizontal line (no dashed underscores)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Access the paragraph XML and add a bottom border
    p_element = p._p
    p_props = p_element.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # 'single' = solid line
    bottom.set(qn('w:sz'), '8')        # Line thickness
    bottom.set(qn('w:space'), '1')     # Padding around line
    bottom.set(qn('w:color'), '000000')
    border.append(bottom)
    p_props.append(border)

def generate_student_worksheet_doc(sections: list) -> str:
    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    title_run = title.add_run("Student Worksheet")
    title_run.font.name = "Poppins"
    title_run.font.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def split_parts(content: str):
        content = content.strip()
        english_part, translation_part = "", ""

        # Split by \n\n — now only need 2 parts
        parts = content.split("\n\n")
        if len(parts) >= 2:
            english_part, translation_part = parts[:2]
        else:
            # fallback to non-ASCII (e.g. Spanish) split
            match = re.search(r'[^\x00-\x7F]', content)
            if match:
                split_index = match.start()
                english_part = content[:split_index].rstrip()
                translation_part = content[split_index:].lstrip()
            else:
                english_part = content

        return english_part.strip(), translation_part.strip()

    # --- Add sections ---
    for section in sections:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section["section"])
        heading_run.font.name = "Poppins SemiBold"
        heading_run.font.size = Pt(18)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        # Split content
        content = section["content"]
        english_part, translation_part = split_parts(content)

        # 🔵 English (question)
        if english_part:
            for line in english_part.split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.font.name = "Poppins"
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 🔴 Translated (question)
        if translation_part:
            for line in translation_part.split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.font.name = "Poppins"
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ✏️ Add 5 blank answer lines
        for _ in range(5):
            add_horizontal_line(doc)

        doc.add_paragraph()  # Spacer

    # --- Save the file ---
    save_dir = "data/outputs/worksheets"
    os.makedirs(save_dir, exist_ok=True)
    filename = f"student_worksheet_{uuid.uuid4().hex}.docx"
    path = os.path.join(save_dir, filename)
    doc.save(path)

    return path